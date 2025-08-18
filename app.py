# app.py
import streamlit as st
import pandas as pd
import numpy as np
import io, csv
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt
from datetime import datetime

# ==============================
# Configuración general
# ==============================
st.set_page_config(page_title="CAAT – Auditoría Automatizada", layout="wide")
st.title("🧪 Herramienta CAAT – Auditoría Automatizada")
st.markdown("Sube archivos y ejecuta las pruebas en cada sección. Soporta **CSV/XLSX/XLS/TXT**.")

# ===== Estilos y helper UI =====
st.markdown("""
<style>
.main .block-container {max-width: 1200px; padding-top: .5rem; padding-bottom: 2rem;}
.section-card {
  border: 1px solid rgba(125,125,125,.25);
  border-radius: 14px; padding: 16px 18px; margin: 18px 0 18px 0;
  background: rgba(200,200,255,.07);
}
.section-title { font-size: 26px; font-weight: 800; margin-bottom: 6px; }
.section-desc  { font-size: 18px; color:#374151; }
.badge {
  display: inline-block; padding: 4px 10px; border-radius: 999px;
  background: #eef2ff; color: #2f3ab2; font-size: 12px; font-weight: 600; margin-left: 6px;
  border: 1px solid rgba(47,58,178,.15);
}
.big-warning { font-size: 16px; line-height: 1.35; }
[data-testid="stFileUploader"] {border-radius: 12px; border: 1px dashed rgba(125,125,125,.35); padding: 14px;}
.stButton>button { border-radius: 999px !important; padding: .55rem 1rem; font-weight: 600; }
</style>
""", unsafe_allow_html=True)

def section_intro(n, title, desc):
    st.markdown(f"""
<div class="section-card">
  <div class="section-title">{n} {title} <span class="badge">CSV/XLSX/XLS/TXT</span></div>
  <div class="section-desc">{desc}</div>
</div>
""", unsafe_allow_html=True)

with st.sidebar:
    st.markdown("### 🧭 Navegación")
    st.write("2) Montos inusuales\n\n3) Conciliación A vs B\n\n4) Benford")
    st.markdown("---")
    st.markdown("### 💡 Consejos")
    st.caption("- Benford: columna de **montos** y muestra **grande**.\n- Conciliación: define **clave** y **tolerancia**.\n- Descargas: resultados en **XLSX** y reportes en **DOCX**.")
    st.markdown("---")
    st.caption("Versión CAAT A-2025 • Streamlit")

# ==============================
# Utilidades de lectura y helpers
# ==============================
SINONIMOS_ID = ["idfactura","id_factura","numero","número","numerofactura","numero_factura",
    "serie","serie_comprobante","clave_acceso","idtransaccion","id_transaccion","referencia","doc","documento","id","idcliente","idproveedor"]
SINONIMOS_MONTO = ["total","monto","importe","valor","monto_total","total_ingresado",
    "importe_total","importe neto","subtotal+iva","total factura","totalfactura","amount","total_amount"]
SINONIMOS_FECHA = ["fecha","fecha_emision","fecha emisión","f_emision","fecha documento",
    "fecha_doc","fechadoc","fecha fact","fecha factura","emision","date","fecha_registro"]

def sniff_delimiter(sample_bytes: bytes):
    try:
        sample = sample_bytes.decode('utf-8', errors='ignore')
        dialect = csv.Sniffer().sniff(sample, delimiters=";,|\t")
        return dialect.delimiter
    except Exception:
        return None

def try_read_csv(file_obj):
    data = file_obj.read()
    if isinstance(data, bytes):
        delim = sniff_delimiter(data[:4096])
        bio = io.BytesIO(data)
        if delim:
            try:
                return pd.read_csv(bio, sep=delim, engine="python")
            except Exception:
                pass
        bio.seek(0)
        try:
            return pd.read_csv(bio, sep=None, engine="python")
        except Exception:
            bio.seek(0)
            return pd.read_csv(bio, sep=None, engine="python", encoding="latin-1")
    else:
        return pd.read_csv(io.StringIO(data))

def try_read_excel(file_obj, widget_key="sheet"):
    xls = pd.ExcelFile(file_obj)
    sheet = st.selectbox("📄 Hoja de Excel", xls.sheet_names, key=widget_key)
    return pd.read_excel(xls, sheet_name=sheet)

def load_any(file, widget_key="sheet"):
    name = file.name.lower()
    if name.endswith(".csv") or name.endswith(".txt"):
        return try_read_csv(file)
    if name.endswith((".xlsx",".xls")):
        return try_read_excel(file, widget_key=widget_key)
    raise ValueError("Formato no soportado")

def normalize_headers(df):
    df.columns = [str(c).strip() for c in df.columns]
    return df

def col_auto(df, candidatos):
    cols_norm = {c.lower().strip(): c for c in df.columns}
    for alias in candidatos:
        if alias in cols_norm:
            return cols_norm[alias]
    for c in df.columns:
        cl = c.lower()
        if any(alias in cl for alias in candidatos):
            return c
    return None

def coerce_amount(series):
    s = series.astype(str)
    s = s.str.replace(r"\.", "", regex=True)  # remover miles con punto
    s = s.str.replace(",", ".", regex=False)  # coma -> punto
    return pd.to_numeric(s, errors="coerce")

def coerce_date(series):
    return pd.to_datetime(series, errors="coerce", dayfirst=True)

def to_xlsx_bytes(df: pd.DataFrame, sheet_name="Hoja1") -> bytes:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name[:31])
    return buffer.getvalue()

def to_multi_xlsx_bytes(sheets: dict) -> bytes:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for name, df in sheets.items():
            safe = str(name)[:31]
            if isinstance(df, pd.DataFrame):
                df.to_excel(writer, index=False, sheet_name=safe)
            else:
                pd.DataFrame({"valor": [df]}).to_excel(writer, index=False, sheet_name=safe)
    return buffer.getvalue()

def docx_from_sections(title: str, sections: list[tuple[str, list[str]]]) -> bytes:
    d = Document()
    d.add_heading(title, level=1)
    for heading, bullets in sections:
        d.add_heading(heading, level=2)
        for item in bullets:
            p = d.add_paragraph(item, style="List Bullet")
            p.style.font.size = Pt(11)
    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()

def pct(x, digits=2):
    try:
        return f"{x*100:.{digits}f}%"
    except Exception:
        return "n/a"

# ======================================================
# 2) CAAT – Detección de Montos Inusuales (mejorado)
# ======================================================
section_intro("2️⃣", "Detección de Montos Inusuales",
              "Encuentra transacciones que superan un umbral (fijo o estadístico) y genera un **reporte con recomendaciones** para el auditor.")

file_unusual = st.file_uploader("📁 Subir archivo para montos inusuales (CSV/XLSX/XLS/TXT)", type=["csv","xlsx","xls","txt"], key="unusual")

if file_unusual:
    try:
        dfm = load_any(file_unusual, widget_key="sheet_unusual")
        dfm = normalize_headers(dfm)
        st.success(f"✅ Archivo cargado. Filas: {len(dfm)}")
        with st.expander("Ver primeras filas"):
            st.dataframe(dfm.head())

        # Selección de columnas
        sugerida_monto = col_auto(dfm, SINONIMOS_MONTO) or (dfm.select_dtypes(include="number").columns.tolist()[:1] or [None])[0]
        col_monto = st.selectbox("💰 Columna de monto", dfm.columns.tolist(),
                                 index=(dfm.columns.tolist().index(sugerida_monto) if sugerida_monto in dfm.columns else 0))
        col_id = st.selectbox("🔑 Columna identificadora (ID/Número/Referencia) (opcional)", ["(ninguna)"] + dfm.columns.tolist(), index=0)
        col_fecha_opt = st.selectbox("📅 Columna de fecha (opcional)", ["(ninguna)"] + dfm.columns.tolist(), index=0)

        metodo = st.radio("Método de detección", ["Umbral fijo", "Umbral estadístico (media + k·σ)"], horizontal=True)
        ejecutar = False
        if metodo.startswith("Umbral fijo"):
            umbral = st.number_input("💵 Umbral fijo ($):", min_value=0.0, value=10000.0)
            ejecutar = st.button("🔍 Ejecutar (fijo)")
        else:
            k = st.slider("🔬 k (media + k·σ)", min_value=1, max_value=5, value=2)
            ejecutar = st.button("🔍 Ejecutar (estadístico)")

        if ejecutar:
            serie_monto = dfm[col_monto]
            if serie_monto.dtype == object:
                serie_monto = coerce_amount(serie_monto)

            dfm["_MONTO_"] = pd.to_numeric(serie_monto, errors="coerce")
            base = dfm.dropna(subset=["_MONTO_"]).copy()

            # Fecha si aplica
            fecha_colname = None
            if col_fecha_opt != "(ninguna)":
                base["_FECHA_"] = coerce_date(base[col_fecha_opt])
                fecha_colname = col_fecha_opt

            if metodo.startswith("Umbral fijo"):
                limite = umbral
                criterio_txt = f"Umbral fijo = {umbral:,.2f}"
            else:
                media = base["_MONTO_"].mean()
                std = base["_MONTO_"].std(ddof=0)
                limite = media + k * std
                criterio_txt = f"Umbral estadístico = media {media:,.2f} + {k}·σ {std:,.2f} → {limite:,.2f}"

            hallazgos = base[base["_MONTO_"] > limite].copy()

            # KPIs
            total_tx = len(base)
            total_h = len(hallazgos)
            prop_h = 0 if total_tx == 0 else total_h / total_tx
            suma_total = base["_MONTO_"].sum()
            suma_h = hallazgos["_MONTO_"].sum()

            st.subheader("📊 Resultados")
            st.write(f"**Criterio aplicado:** {criterio_txt}")
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Transacciones analizadas", total_tx)
            c2.metric("Hallazgos (n)", total_h)
            c3.metric("% del total", f"{prop_h*100:.2f}%")
            c4.metric("Suma hallazgos", f"{suma_h:,.2f}")

            # Enriquecimiento: z-score y tops
            if total_h > 0:
                mu = base["_MONTO_"].mean()
                sd = base["_MONTO_"].std(ddof=0) or 1.0
                hallazgos["_zscore_"] = (hallazgos["_MONTO_"] - mu) / sd

                top_monto = hallazgos.sort_values("_MONTO_", ascending=False).head(20)
                top_z = hallazgos.sort_values("_zscore_", ascending=False).head(20)

                # Agrupación por ID si aplica
                grp_df = pd.DataFrame()
                if col_id != "(ninguna)" and col_id in hallazgos.columns:
                    grp_df = (hallazgos.groupby(col_id, dropna=False)
                                        .agg(N=(" _MONTO_".strip(),"count"),
                                             Suma=("_MONTO_","sum"),
                                             Max=("_MONTO_","max"))
                                        .sort_values("Suma", ascending=False)
                                        .head(20)
                              )

                # Mostrar y descargar XLSX multi-hojas
                sheets = {"Hallazgos": hallazgos}
                stats = pd.DataFrame({
                    "Métrica":["Total tx","Hallazgos","% hallazgos","Suma total","Suma hallazgos","Criterio"],
                    "Valor":[total_tx, total_h, f"{prop_h*100:.2f}%", f"{suma_total:,.2f}", f"{suma_h:,.2f}", criterio_txt]
                })
                sheets["ResumenEstadistico"] = stats
                sheets["TopPorMonto"] = top_monto
                sheets["TopPorZscore"] = top_z
                if not grp_df.empty:
                    sheets["GrupoPorID"] = grp_df.reset_index()

                xlsx_bytes = to_multi_xlsx_bytes(sheets)
                st.download_button("⬇️ Descargar hallazgos y resúmenes (XLSX)",
                                   xlsx_bytes, "montos_inusuales.xlsx",
                                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

                st.dataframe(hallazgos.head(1000))

                # Reporte DOCX detallado
                fecha_rango = "Sin fecha" if fecha_colname is None else f"{base['_FECHA_'].min()} → {base['_FECHA_'].max()}"
                bullets_resumen = [
                    f"Archivo analizado: {file_unusual.name}",
                    f"Filas válidas: {total_tx}",
                    f"Hallazgos: {total_h} ({prop_h*100:.2f}%)",
                    f"Suma hallazgos: {suma_h:,.2f} (vs total {suma_total:,.2f})",
                    f"Criterio aplicado: {criterio_txt}",
                    f"Rango de fechas: {fecha_rango}"
                ]

                # Detalle basado en contenido
                detalle = []
                if total_h > 0:
                    may = hallazgos["_MONTO_"].max()
                    prom_h = hallazgos["_MONTO_"].mean()
                    detalle += [
                        f"Mayor hallazgo: {may:,.2f}",
                        f"Promedio de hallazgos: {prom_h:,.2f}"
                    ]
                    if col_id != "(ninguna)" and col_id in hallazgos.columns:
                        top_id = (hallazgos.groupby(col_id)["_MONTO_"].sum().sort_values(ascending=False).head(5))
                        detalle.append("Top 5 por suma (ID): " + ", ".join([f"{idx}: {val:,.2f}" for idx, val in top_id.items()]))

                recomendaciones = [
                    "Solicitar respaldo documental (OC, contratos, aprobaciones) para cada hallazgo y validar trazabilidad en el sistema.",
                    "Aplicar revisiones dirigidas sobre los 20 mayores importes y los 20 mayores z-scores.",
                    "Verificar límites de aprobación y segregación de funciones; contrastar con el flujo de autorizaciones.",
                    "Investigar concentración por ID (cliente/proveedor) si corresponde; evaluar riesgo de fraude o sobre-facturación.",
                    "Cruzar con políticas de precios y descuentos; confirmar cálculo de impuestos y redondeos.",
                    "Si hay patrón por fechas (cierres, fines de mes), revisar asientos manuales y notas de ajuste.",
                    "Ampliar muestra si el % de hallazgos supera el umbral de materialidad definido por auditoría."
                ]

                sections = [
                    ("RESUMEN", [f"• {x}" for x in bullets_resumen]),
                    ("DETALLE PRINCIPAL", [f"• {x}" for x in detalle] if detalle else ["• No se encontraron detalles adicionales."]),
                    ("RECOMENDACIONES", [f"• {x}" for x in recomendaciones]),
                    ("REFERENCIA XLSX", ["• Los hallazgos, tops y agrupaciones se incluyen en el archivo 'montos_inusuales.xlsx' adjunto."])
                ]
                st.download_button("⬇️ Descargar reporte detallado (DOCX)",
                                   docx_from_sections("Montos Inusuales – Reporte de Auditoría", sections),
                                   "reporte_montos_inusuales.docx",
                                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.success("✅ No se encontraron montos inusuales con el criterio seleccionado.")

    except Exception as e:
        st.error(f"❌ Error: {e}")

# ======================================================
# 3) CAAT – Conciliación de Reportes (A vs. B) + recomendaciones
# ======================================================
section_intro("3️⃣", "Conciliación de Reportes (A vs. B)",
              "Compara dos archivos (p. ej., facturación y contabilidad) y genera un **informe con hallazgos (XLSX) y recomendaciones (DOCX)**.")

colA, colB = st.columns(2)
with colA:
    file_A = st.file_uploader("📁 Archivo A (p.ej., Facturación)", type=["csv","xlsx","xls","txt"], key="conc_a")
with colB:
    file_B = st.file_uploader("📁 Archivo B (p.ej., Contabilidad)", type=["csv","xlsx","xls","txt"], key="conc_b")

if file_A and file_B:
    try:
        A = load_any(file_A, widget_key="sheet_A"); A = normalize_headers(A)
        B = load_any(file_B, widget_key="sheet_B"); B = normalize_headers(B)

        st.success(f"✅ Cargados A={len(A)} filas, B={len(B)} filas")
        with st.expander("Ver primeras filas"):
            st.write("A (preview)"); st.dataframe(A.head())
            st.write("B (preview)"); st.dataframe(B.head())

        comunes = [c for c in A.columns if c in set(B.columns)]
        if not comunes:
            st.error("❌ No hay columnas en común entre A y B.")
            st.stop()

        clave_sug = col_auto(A, SINONIMOS_ID) if col_auto(A, SINONIMOS_ID) in comunes else comunes[0]
        montoA_sug = col_auto(A, SINONIMOS_MONTO) or (A.select_dtypes(include="number").columns.tolist()[:1] or [None])[0]
        montoB_sug = col_auto(B, SINONIMOS_MONTO) or (B.select_dtypes(include="number").columns.tolist()[:1] or [None])[0]
        fechaA_sug = col_auto(A, SINONIMOS_FECHA)
        fechaB_sug = col_auto(B, SINONIMOS_FECHA)

        st.subheader("🔧 Configuración")
        clave = st.selectbox("🔑 Columna clave común", comunes, index=(comunes.index(clave_sug) if (clave_sug in comunes) else 0))
        monto_A = st.selectbox("💰 Columna de monto en A", A.columns.tolist(), index=(A.columns.tolist().index(montoA_sug) if (montoA_sug in A.columns) else 0))
        monto_B = st.selectbox("💰 Columna de monto en B", B.columns.tolist(), index=(B.columns.tolist().index(montoB_sug) if (montoB_sug in B.columns) else 0))
        fecha_A_opt = st.selectbox("📅 Columna de fecha en A (opcional)", ["(ninguna)"] + A.columns.tolist(),
                                   index=(["(ninguna)"] + A.columns.tolist()).index(fechaA_sug) if (fechaA_sug in A.columns) else 0)
        fecha_B_opt = st.selectbox("📅 Columna de fecha en B (opcional)", ["(ninguna)"] + B.columns.tolist(),
                                   index=(["(ninguna)"] + B.columns.tolist()).index(fechaB_sug) if (fechaB_sug in B.columns) else 0)
        tolerancia = st.number_input("🎯 Tolerancia para diferencias de monto (valor absoluto)", min_value=0.0, value=0.0)

        if st.button("🔍 Ejecutar conciliación"):
            # Normalización
            A["_CLAVE_"] = A[clave].astype(str).str.strip().str.upper()
            B["_CLAVE_"] = B[clave].astype(str).str.strip().str.upper()
            A["_MONTO_"] = coerce_amount(A[monto_A]) if A[monto_A].dtype == object else pd.to_numeric(A[monto_A], errors="coerce")
            B["_MONTO_"] = coerce_amount(B[monto_B]) if B[monto_B].dtype == object else pd.to_numeric(B[monto_B], errors="coerce")
            if fecha_A_opt != "(ninguna)": A["_FECHA_"] = coerce_date(A[fecha_A_opt])
            if fecha_B_opt != "(ninguna)": B["_FECHA_"] = coerce_date(B[fecha_B_opt])

            merged = A.merge(B, on="_CLAVE_", how="outer", suffixes=("_A","_B"), indicator=True)
            solo_A = merged[merged["_merge"]=="left_only"].copy()
            solo_B = merged[merged["_merge"]=="right_only"].copy()
            coinc   = merged[merged["_merge"]=="both"].copy()
            coinc["_diff_monto"] = (coinc["_MONTO__A"] - coinc["_MONTO__B"])
            coinc["_diff_monto_abs"] = coinc["_diff_monto"].abs()
            diff_monto = coinc[coinc["_diff_monto_abs"] > tolerancia].copy()

            diff_fecha = pd.DataFrame()
            if "_FECHA__A" in coinc.columns and "_FECHA__B" in coinc.columns:
                diff_fecha = coinc[(~coinc["_FECHA__A"].isna()) & (~coinc["_FECHA__B"].isna()) & (coinc["_FECHA__A"] != coinc["_FECHA__B"])][["_CLAVE_","_FECHA__A","_FECHA__B"]].copy()

            # KPIs y totales
            total_A = A["_MONTO_"].sum(skipna=True)
            total_B = B["_MONTO_"].sum(skipna=True)
            delta_total = total_A - total_B
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Solo en A", len(solo_A)); c2.metric("Solo en B", len(solo_B))
            c3.metric("Dif. de monto (> tol.)", len(diff_monto)); c4.metric("Δ Total A-B", f"{delta_total:,.2f}")

            # Mostrar y descargar XLSX consolidado
            sheets_conc = {
                "Resumen": pd.DataFrame({
                    "Métrica": ["Total A","Total B","Δ A-B","Solo en A (n)","Solo en B (n)","Dif. monto (n)","Dif. fecha (n)","Tolerancia"],
                    "Valor":   [f"{total_A:,.2f}", f"{total_B:,.2f}", f"{delta_total:,.2f}", len(solo_A), len(solo_B), len(diff_monto), len(diff_fecha), f"{tolerancia:,.2f}"]
                }),
                "Solo_en_A": solo_A,
                "Solo_en_B": solo_B,
                "Diferencias_Monto": diff_monto[["_CLAVE_","_MONTO__A","_MONTO__B","_diff_monto","_diff_monto_abs"]],
                "Diferencias_Fecha": diff_fecha if not diff_fecha.empty else pd.DataFrame(columns=["_CLAVE_","_FECHA__A","_FECHA__B"])
            }
            st.download_button("⬇️ Descargar hallazgos (XLSX)", to_multi_xlsx_bytes(sheets_conc),
                               "hallazgos_conciliacion.xlsx",
                               "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

            # Mostrar vistas resumidas
            with st.expander("🟦 Solo en A"):
                st.dataframe(solo_A.head(1000))
            with st.expander("🟧 Solo en B"):
                st.dataframe(solo_B.head(1000))
            with st.expander("🟥 Coincidentes con diferencias de monto"):
                st.dataframe(diff_monto[["_CLAVE_","_MONTO__A","_MONTO__B","_diff_monto","_diff_monto_abs"]].head(1000))
            if not diff_fecha.empty:
                with st.expander("🟨 Coincidentes con diferencias de fecha"):
                    st.dataframe(diff_fecha.head(1000))

            # Reporte DOCX profundo
            # Top 10 diferencias por monto absoluto
            top_dif = diff_monto.sort_values("_diff_monto_abs", ascending=False).head(10)
            pos = diff_monto[diff_monto["_diff_monto"] > 0]["_diff_monto"].sum()
            neg = diff_monto[diff_monto["_diff_monto"] < 0]["_diff_monto"].sum()
            bullets_resumen = [
                f"Archivo A: {file_A.name} | Archivo B: {file_B.name}",
                f"Total A: {total_A:,.2f} | Total B: {total_B:,.2f} | Δ A-B: {delta_total:,.2f}",
                f"Solo en A: {len(solo_A)} | Solo en B: {len(solo_B)}",
                f"Dif. de monto (> tolerancia): {len(diff_monto)} | Suma dif. positivas: {pos:,.2f} | negativas: {neg:,.2f}",
                f"Dif. de fecha: {len(diff_fecha)}",
                f"Tolerancia aplicada: {tolerancia:,.2f}"
            ]

            detalle = []
            for i, row in top_dif.iterrows():
                detalle.append(f"{row.get('_CLAVE_', 's/clave')}: A={row.get('_MONTO__A',np.nan):,.2f} | B={row.get('_MONTO__B',np.nan):,.2f} | Δ={row.get('_diff_monto',np.nan):,.2f} (|Δ|={row.get('_diff_monto_abs',np.nan):,.2f})")
            if not detalle:
                detalle = ["No hay diferencias de monto sobre la tolerancia."]

            recomendaciones = [
                "Revisar interfaz/integración entre sistemas (logs de carga, horarios de corte, reprocesos).",
                "Para diferencias de monto: confirmar tipo de cambio, descuentos, impuestos, notas de crédito y redondeos.",
                "Validar que no existan asientos manuales fuera del flujo autorizado; revisar bitácoras y perfiles.",
                "Programar conciliaciones periódicas automáticas con alarmas por materialidad.",
                "Investigar registros Solo en A/B: reprocesar interfaz y verificar dependencia temporal (fechas cercanas a cierres).",
                "Acordar con dueños de procesos un marco de tolerancia por tipo de transacción.",
            ]

            sections = [
                ("RESUMEN", [f"• {x}" for x in bullets_resumen]),
                ("TOP 10 DIFERENCIAS POR MONTO", [f"• {x}" for x in detalle]),
                ("RECOMENDACIONES", [f"• {x}" for x in recomendaciones]),
                ("REFERENCIA XLSX", ["• Ver 'hallazgos_conciliacion.xlsx' con hojas de Solo_en_A, Solo_en_B, Diferencias_Monto, Diferencias_Fecha y Resumen."])
            ]
            st.download_button("⬇️ Descargar reporte detallado (DOCX)",
                               docx_from_sections("Conciliación A vs. B – Reporte de Auditoría", sections),
                               "reporte_conciliacion.docx",
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        st.error(f"❌ Error en conciliación: {e}")

# ======================================================
# 4) CAAT – Ley de Benford (con sospechosos + XLSX + DOCX)
# ======================================================
section_intro("4️⃣", "Ley de Benford aplicada a transacciones",
              "Contrasta el primer dígito de los montos con la distribución esperada por Benford, **lista transacciones sospechosas** y emite un **reporte**.")

st.markdown("""
<div class="section-card">
<div class="big-warning">
<strong>⚠️ Advertencia importante:</strong> La Ley de Benford es adecuada para conjuntos grandes de datos
de naturaleza espontánea (no pre-condicionados). No usar con series acotadas, precios fijos,
mínimos/máximos impuestos, folios o montos prefijados; los resultados podrían ser engañosos.
</div>
</div>
""", unsafe_allow_html=True)

file_benford = st.file_uploader("📁 Subir archivo (CSV/XLSX/XLS/TXT)", type=["csv","xlsx","xls","txt"], key="benford")

def first_digit_series(series: pd.Series) -> pd.Series:
    if series.dtype == object: x = coerce_amount(series)
    else: x = pd.to_numeric(series, errors="coerce")
    x = x.abs(); x = x[x > 0].dropna()
    s = x.apply(lambda v: f"{v:.15g}")
    s = s.str.replace(".", "", regex=False).str.lstrip("0")
    first = s.str[0].dropna()
    first = first[first.str.contains(r"[1-9]", regex=True)]
    return first.astype(int)

def benford_expected() -> pd.Series:
    d = np.arange(1, 10); p = np.log10(1 + 1/d)
    return pd.Series(p, index=d)

if file_benford:
    try:
        dfb = load_any(file_benford, widget_key="sheet_benford"); dfb = normalize_headers(dfb)
        st.success(f"✅ Archivo cargado. Filas: {len(dfb)}")
        with st.expander("Ver primeras filas"): st.dataframe(dfb.head())

        # Columnas candidatas a monto
        def is_amount_candidate(s: pd.Series) -> bool:
            if pd.api.types.is_numeric_dtype(s): return True
            if s.dtype == object:
                conv = coerce_amount(s); return conv.notna().mean() >= 0.30
            return False
        candidatas = [c for c in dfb.columns if is_amount_candidate(dfb[c])]
        if not candidatas:
            st.error("No se hallaron columnas de monto válidas. Debe existir una columna numérica o convertible.")
            st.stop()

        sugerida_monto_b = col_auto(dfb[candidatas], SINONIMOS_MONTO) or candidatas[0]
        col_monto_b = st.selectbox("💰 Columna de monto", candidatas, index=(candidatas.index(sugerida_monto_b) if sugerida_monto_b in candidatas else 0))
        min_val = st.number_input("🔻 Ignorar montos menores a (opcional)", min_value=0.0, value=0.0)
        min_count_alert = st.number_input("🔔 Mínimo sugerido de observaciones", min_value=0, value=100)
        desvio_min = st.number_input("🎚 Umbral de desviación por dígito (puntos porcentuales)", min_value=0.0, value=2.0, step=0.5,
                                     help="Se marcan 'sospechosos' los dígitos cuya (Observado - Esperado) ≥ este umbral.")

        if st.button("🔍 Ejecutar Benford"):
            serie_orig = dfb[col_monto_b]
            if serie_orig.dtype == object:
                conv = coerce_amount(serie_orig)
                if conv.notna().mean() < 0.30:
                    st.error("La columna seleccionada no tiene suficientes valores convertibles a número (≥30%).")
                    st.stop()
                serie_num = conv
            else:
                serie_num = pd.to_numeric(serie_orig, errors="coerce")

            serie_num = serie_num.dropna()
            if min_val > 0:
                serie_num = serie_num[serie_num.abs() >= min_val]

            n_total = len(serie_num)
            fd = first_digit_series(serie_num)  # <-- FIX: pasar Series, no ndarray
            n = len(fd)
            if n == 0:
                st.error("❌ No hay suficientes datos numéricos válidos tras la limpieza/filtros.")
            else:
                obs_counts = fd.value_counts().reindex(range(1,10), fill_value=0).sort_index()
                obs_prop   = obs_counts / n
                exp_prop   = benford_expected()
                exp_counts = (exp_prop * n)
                chi2 = (((obs_counts - exp_counts) ** 2) / exp_counts.replace(0, np.nan)).sum()
                chi2_crit = 15.507  # α=0.05, gl=8
                cumple = chi2 <= chi2_crit

                st.subheader("📊 Resumen Benford")
                c1, c2, c3 = st.columns(3)
                c1.metric("Observaciones válidas", n); c2.metric("Chi-cuadrado", f"{chi2:,.3f}"); c3.metric("¿Cumple (α=0.05)?", "Sí ✅" if cumple else "No ⚠️")
                if n < min_count_alert: st.info(f"ℹ️ Nota: {n} observaciones; sugerido ≥ {min_count_alert} para mayor robustez.")

                # Tabla observados vs esperados
                tabla = pd.DataFrame({
                    "Dígito": list(range(1,10)),
                    "Frecuencia Observada": obs_counts.values,
                    "Proporción Observada (%)": (obs_prop.values * 100).round(2),
                    "Proporción Esperada (%)": (exp_prop.values * 100).round(2),
                    "Desviación (pp)": ((obs_prop - exp_prop).values * 100).round(2)
                })
                st.dataframe(tabla)

                # Gráfico barras
                fig, ax = plt.subplots()
                idx = np.arange(1, 10)
                ax.bar(idx - 0.15, obs_prop.values, width=0.3, label="Observado")
                ax.bar(idx + 0.15, exp_prop.values, width=0.3, label="Esperado (Benford)")
                ax.set_xticks(idx); ax.set_xlabel("Primer dígito"); ax.set_ylabel("Proporción")
                ax.set_title("Ley de Benford: Observado vs. Esperado"); ax.legend()
                st.pyplot(fig)

                # Dígitos sospechosos y transacciones asociadas (con filas originales)
                sospechosos = tabla.loc[tabla["Desviación (pp)"] >= desvio_min, "Dígito"].tolist()
                st.write(f"**Dígitos 'sospechosos' (desviación ≥ {desvio_min:.1f} pp):** {sospechosos if sospechosos else 'Ninguno'}")

                # Mapear primer dígito a cada fila válida y quedarse con las sospechosas
                first_digits_full = first_digit_series(serie_num)
                sospechosas_idx = first_digits_full[first_digits_full.isin(sospechosos)].index
                sospechosas_rows = dfb.loc[sospechosas_idx].copy()
                sospechosas_rows["_monto_convertido_"] = serie_num.loc[sospechosas_idx]
                sospechosas_rows["_1er_dig"] = first_digits_full.loc[sospechosas_idx]

                with st.expander("🔎 Ver transacciones sospechosas por dígito"):
                    if not sospechosas_rows.empty:
                        st.dataframe(sospechosas_rows.head(1000))
                        st.download_button("⬇️ Descargar sospechosas (XLSX)",
                                           to_xlsx_bytes(sospechosas_rows, "Sospechosas_Benford"),
                                           "benford_sospechosas.xlsx",
                                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                    else:
                        st.info("No se encontraron transacciones sospechosas con el umbral seleccionado.")

                # Descarga XLSX del resumen
                st.download_button("⬇️ Descargar tabla resumen (XLSX)",
                                   to_xlsx_bytes(tabla, "Resumen_Benford"),
                                   "benford_resumen.xlsx",
                                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

                # Reporte DOCX detallado
                recomendaciones = [
                    "Confirmar que el conjunto es adecuado para Benford (no pre-condicionado; suficiente volumen; diversidad de magnitudes).",
                    "Para los dígitos con mayor desviación, revisar una muestra dirigida de transacciones (comprobantes, aprobaciones).",
                    "Analizar por segmentos (por ejemplo, por proveedor, centro de costo, periodo) para identificar focos específicos.",
                    "Investigar reglas de redondeo, precios fijos, topes o mínimos que puedan explicar desviaciones.",
                    "Si persisten desviaciones materiales sin causa operativa, elevar como indicio y aplicar pruebas forenses complementarias."
                ]
                resumen = [
                    f"Archivo: {file_benford.name}",
                    f"Observaciones válidas: {n} (de {n_total} originales tras filtros)",
                    f"Chi-cuadrado: {chi2:,.3f} | Criterio α=0.05, gl=8 → {'Cumple' if cumple else 'No cumple'}",
                    "Dígitos con mayor desviación (pp): " + ", ".join([f"{int(d)}" for d in sospechosos]) if sospechosos else "Sin dígitos con desviación sobre el umbral",
                    f"Umbral de desviación usado: {desvio_min:.1f} pp",
                    f"Filtro 'ignorar menores a': {min_val:,.2f}"
                ]
                sections = [
                    ("RESUMEN", [f"• {x}" for x in resumen]),
                    ("DESVIACIONES POR DÍGITO", [f"• {row.Dígito}: Obs {row['Proporción Observada (%)']}% vs Exp {row['Proporción Esperada (%)']}% (Δ {row['Desviación (pp)']} pp)" for _, row in tabla.iterrows()]),
                    ("RECOMENDACIONES", [f"• {x}" for x in recomendaciones]),
                    ("REFERENCIA XLSX", [
                        "• 'benford_resumen.xlsx' incluye tabla de observados vs esperados.",
                        "• 'benford_sospechosas.xlsx' lista las transacciones cuyas cifras iniciales pertenecen a dígitos con desviación sobre el umbral."
                    ])
                ]
                st.download_button("⬇️ Descargar reporte detallado (DOCX)",
                                   docx_from_sections("Ley de Benford – Reporte de Auditoría", sections),
                                   "reporte_benford.docx",
                                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        st.error(f"❌ Error en Benford: {e}")
